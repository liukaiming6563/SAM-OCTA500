r"""
生成 OCTA-500 数据集处理结果 DOCX 报告。

命令示例：
    python scripts/data_prepare/08_generate_octa500_dataset_report_docx.py --config configs/data/octa500_prepare.yaml

输入：
    本地数据目录中的 raw 扫描报告、全任务 processed 检查摘要、manifest 和目录结构。

输出：
    D:\study\project\Data\DataSet\OCTA500\OCTA500_processed_dataset_report.docx

报告内容：
    1. 数据目录总览。
    2. zip_original、raw、processed、outputs 四类目录含义。
    3. 3M/6M 与五类分割任务的样本数量、split、检查结果。
    4. 对后续深度学习任务的作用，包括 Dataset、训练、验证、测试和可视化。
"""

from __future__ import annotations

import argparse
import json
import sys
from datetime import datetime
from pathlib import Path
from typing import Any

PROJECT_ROOT = Path(__file__).resolve().parents[2]
sys.path.insert(0, str(PROJECT_ROOT / "src"))

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

from sam_octa500.utils.config_utils import load_config
from sam_octa500.utils.log_utils import log_header, write_log
from sam_octa500.utils.path_utils import as_path


def parse_args() -> argparse.Namespace:
    """解析命令行参数。"""
    parser = argparse.ArgumentParser(description="生成 OCTA-500 processed 数据集 DOCX 报告")
    parser.add_argument("--config", required=True, help="数据准备配置文件路径")
    parser.add_argument("--output", default="", help="可选输出 DOCX 路径")
    return parser.parse_args()


def load_json(path: Path) -> dict[str, Any]:
    """
    读取 JSON 文件。

    参数：
        path: JSON 文件路径。

    返回：
        JSON 字典；文件不存在时返回空字典。
    """
    if not path.exists():
        return {}
    return json.loads(path.read_text(encoding="utf-8"))


def count_files(path: Path, pattern: str = "*") -> int:
    """
    统计目录下文件数量。

    参数：
        path: 待统计目录。
        pattern: glob 模式。

    返回：
        文件数量；目录不存在时返回 0。
    """
    if not path.exists():
        return 0
    return sum(1 for p in path.rglob(pattern) if p.is_file())


def setup_styles(document: Document) -> None:
    """
    设置 Word 文档基础样式。

    参数：
        document: python-docx 文档对象。
    """
    section = document.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)

    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.15

    for name, size, color in [
        ("Title", 22, RGBColor(11, 37, 69)),
        ("Heading 1", 16, RGBColor(46, 116, 181)),
        ("Heading 2", 13, RGBColor(31, 77, 120)),
        ("Heading 3", 11.5, RGBColor(31, 77, 120)),
    ]:
        style = styles[name]
        style.font.name = "Microsoft YaHei"
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(8)
        style.paragraph_format.space_after = Pt(5)


def add_kv_table(document: Document, rows: list[tuple[str, str]]) -> None:
    """
    添加两列表格，用于展示路径和说明。

    参数：
        document: Word 文档对象。
        rows: (字段, 内容) 列表。
    """
    table = document.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.autofit = False
    table.columns[0].width = Inches(1.65)
    table.columns[1].width = Inches(4.95)
    headers = table.rows[0].cells
    headers[0].text = "项目"
    headers[1].text = "说明"
    for cell in headers:
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
    for key, value in rows:
        cells = table.add_row().cells
        cells[0].text = key
        cells[1].text = value
        for cell in cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_task_table(document: Document, tasks: list[dict[str, Any]]) -> None:
    """
    添加 processed 全任务统计表。

    参数：
        document: Word 文档对象。
        tasks: processed_all_tasks_summary.json 中的任务列表。
    """
    headers = ["FOV", "任务", "train", "val", "test", "总数", "异常", "前景比例均值"]
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.autofit = True
    for index, header in enumerate(headers):
        table.rows[0].cells[index].text = header
    for task in tasks:
        cells = table.add_row().cells
        cells[0].text = str(task["fov"])
        cells[1].text = str(task["task"])
        cells[2].text = str(task["train"])
        cells[3].text = str(task["val"])
        cells[4].text = str(task["test"])
        cells[5].text = str(task["total"])
        cells[6].text = str(task["anomaly_count"])
        cells[7].text = f"{task.get('foreground_ratio', {}).get('mean', 0):.6f}"
    for row in table.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER


def add_bullet(document: Document, text: str) -> None:
    """
    添加项目符号段落。

    参数：
        document: Word 文档对象。
        text: 项目符号文本。
    """
    document.add_paragraph(text, style="List Bullet")


def build_report(config: dict[str, Any], output_path: Path) -> None:
    """
    构建 DOCX 报告。

    参数：
        config: 数据准备配置。
        output_path: DOCX 输出路径。
    """
    data_root = as_path(config, "data_root")
    raw_root = as_path(config, "raw_root")
    processed_root = as_path(config, "processed_root")
    output_root = as_path(config, "output_root")
    report_dir = as_path(config, "report_dir")
    manifest_dir = as_path(config, "manifest_dir")
    visualization_dir = as_path(config, "visualization_dir")

    raw_report = load_json(report_dir / "raw_structure_report.json")
    all_summary = load_json(report_dir / "processed_all_tasks_summary.json")
    tasks = all_summary.get("tasks", [])
    total_images = sum(int(t.get("total", 0)) for t in tasks)
    total_masks = total_images
    total_anomalies = sum(int(t.get("anomaly_count", 0)) for t in tasks)

    document = Document()
    setup_styles(document)
    document.add_heading("OCTA-500 数据集处理结果报告", level=0)
    p = document.add_paragraph()
    p.add_run("报告生成时间：").bold = True
    p.add_run(datetime.now().strftime("%Y-%m-%d %H:%M:%S"))
    p = document.add_paragraph()
    p.add_run("数据根目录：").bold = True
    p.add_run(str(data_root))

    document.add_heading("1. 处理结果概览", level=1)
    document.add_paragraph(
        "本报告说明 OCTA-500 数据集在本地经过解压、扫描、二维转换、质量检查和可视化抽样后形成的目录结构与文件含义。"
        "processed 目录已经整理为标准 image/mask 分割格式，可直接服务于 PyTorch、SAM、MedSAM、SAM-OCTA、U-Net 等深度学习任务。"
    )
    add_kv_table(
        document,
        [
            ("raw 文件总数", str(raw_report.get("file_count", 0))),
            ("3M 样本", "10301-10500，共 200 个样本"),
            ("6M 样本", "10001-10300，共 300 个样本"),
            ("processed 任务数", f"{len(tasks)} 个：2 个视场 × 5 类标签"),
            ("processed 图像数", f"{total_images} 张 image PNG"),
            ("processed 标签数", f"{total_masks} 张 mask PNG"),
            ("检查异常数", str(total_anomalies)),
            ("可视化抽样", f"{len(tasks)} 个任务，每个任务 20 张，共 {len(tasks) * 20} 张抽样图"),
        ],
    )

    document.add_heading("2. 根目录内容说明", level=1)
    add_kv_table(
        document,
        [
            ("zip_original", "用户规划的原始压缩包目录。本次实际提供的 zip 文件位于数据根目录，未移动、未删除、未修改；该目录保留用于后续规范存放。"),
            ("raw", "解压后的原始数据目录，只读使用。包含 Code、Label、OCTA_3mm、OCTA_6mm，不应在训练或转换过程中直接修改。"),
            ("processed", "标准化后的二维 image/mask 数据目录。每个任务都按 train、val、test 拆分，适合深度学习训练和评估。"),
            ("outputs", "运行产物目录，保存 manifest、日志、检查报告、配置快照和可视化抽样图，用于复现与质量追踪。"),
        ],
    )

    document.add_heading("3. raw 目录说明", level=1)
    document.add_paragraph("raw 目录保存原始解压内容，作为后续所有转换的来源。转换脚本只读取 raw，不会覆盖或修改 raw。")
    add_kv_table(
        document,
        [
            ("raw\\Code", "数据集官方代码或说明文件。可用于追踪官方划分、格式或处理参考。"),
            ("raw\\Label", "原始标签目录，包含 GT_FAZ、GT_LargeVessel、GT_Capillary、GT_Artery、GT_Vein 等标签。"),
            ("raw\\OCTA_3mm", "3mm 视场原始 OCT/OCTA 数据，样本编号 10301-10500。"),
            ("raw\\OCTA_6mm", "6mm 视场原始 OCT/OCTA 数据，样本编号 10001-10300。"),
        ],
    )
    document.add_paragraph("raw 扫描得到的主要后缀分布如下。")
    suffix_rows = [(k, str(v)) for k, v in raw_report.get("suffix_counts", {}).items()]
    add_kv_table(document, suffix_rows if suffix_rows else [("无", "未读取到后缀统计")])

    document.add_heading("4. processed 目录说明", level=1)
    document.add_paragraph(
        "processed\\OCTA500_2D 是本阶段最重要的训练数据目录。它把原始 OCTA(FULL) 投影图和各类标签转换为统一的 PNG 文件，"
        "并按视场、任务和 train/val/test 拆分。"
    )
    add_task_table(document, tasks)

    document.add_heading("4.1 单个任务目录结构", level=2)
    document.add_paragraph("每个任务目录的结构一致，例如 processed\\OCTA500_2D\\3M\\FAZ：")
    add_kv_table(
        document,
        [
            ("train\\images", "训练集输入图像，RGB PNG。用于模型参数学习。"),
            ("train\\masks", "训练集标签，单通道 PNG，背景为 0，前景为 255。用于计算分割损失。"),
            ("val\\images", "验证集输入图像。用于调参、早停和模型选择，不参与参数更新。"),
            ("val\\masks", "验证集标签。用于验证 Dice、IoU、Precision、Recall 等指标。"),
            ("test\\images", "测试集输入图像。用于最终泛化性能评估。"),
            ("test\\masks", "测试集标签。用于最终测试指标计算。"),
            ("meta.csv", "样本级索引表，记录 sample_id、split、原始路径、输出路径、图像和标签尺寸。"),
        ],
    )

    document.add_heading("5. outputs 目录说明", level=1)
    add_kv_table(
        document,
        [
            ("outputs\\manifest", f"保存 raw 和 processed 的 CSV 索引。当前任务级 processed manifest 数量：{count_files(manifest_dir, 'manifest_processed_*.csv')}。"),
            ("outputs\\logs", "保存解压、扫描、转换、检查、可视化、文档生成和批量转换日志。用于排查缺失样本或重复候选。"),
            ("outputs\\reports", "保存 JSON 和 Markdown 检查报告，包括 raw 结构报告、单任务检查报告和全任务摘要。"),
            ("outputs\\visualizations", f"保存抽样可视化图。processed 每个任务 20 张 image/mask/overlay 三联图，便于人工确认标签位置和质量。"),
            ("outputs\\config_snapshots", "保存每次转换时的配置文件副本，保证后续可以追溯当时使用的路径、关键词和 split 规则。"),
        ],
    )

    document.add_heading("6. 对深度学习任务的作用", level=1)
    add_bullet(document, "统一 image/mask 格式：训练脚本无需再理解 OCTA-500 原始复杂目录，只需要读取 processed 下的 PNG 和 meta.csv。")
    add_bullet(document, "固定 train/val/test：避免实验时数据泄漏，便于不同模型之间公平比较。")
    add_bullet(document, "多任务标签齐全：FAZ、大血管、毛细血管、动脉、静脉可分别训练单任务模型，也可扩展为多任务或多类别分割。")
    add_bullet(document, "质量检查报告：提前发现尺寸不一致、mask 非二值、全黑标签、image/mask 文件名不匹配等问题。")
    add_bullet(document, "可视化抽样：为论文实验、数据审核和模型错误分析提供直观依据。")
    add_bullet(document, "配置快照与日志：保证同一批数据可以被复现，后续更换模型或损失函数时数据版本保持稳定。")

    document.add_heading("7. 建议使用方式", level=1)
    document.add_paragraph("后续训练建议直接使用代码工程中的 OCTA500Dataset 读取 processed/OCTA500_2D：")
    add_bullet(document, "输入图像 image 输出形状为 [3, H, W]，适合 CNN、Transformer、SAM 类模型。")
    add_bullet(document, "标签 mask 输出形状为 [1, H, W]，适合二值分割 loss 和 Dice/IoU 指标。")
    add_bullet(document, "可以先从 3M+FAZ 调试训练流程，再扩展到 6M 或血管类别任务。")
    add_bullet(document, "如果进行 SAM/MedSAM 训练，可用 mask 自动生成 bbox、point prompt 或区域提示。")

    document.add_heading("8. 关键路径清单", level=1)
    add_kv_table(
        document,
        [
            ("数据根目录", str(data_root)),
            ("raw 根目录", str(raw_root)),
            ("processed 根目录", str(processed_root / "OCTA500_2D")),
            ("全任务摘要", str(report_dir / "processed_all_tasks_summary.json")),
            ("可视化目录", str(visualization_dir / "processed_samples")),
            ("本 DOCX 报告", str(output_path)),
        ],
    )

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)


def main() -> None:
    """主函数：读取配置并生成 DOCX 报告。"""
    args = parse_args()
    config = load_config(args.config)
    data_root = as_path(config, "data_root")
    output_path = Path(args.output) if args.output else data_root / "OCTA500_processed_dataset_report.docx"
    log_path = as_path(config, "log_dir") / "dataset_report_docx_log.txt"
    lines = log_header("08_generate_octa500_dataset_report_docx.py")
    build_report(config, output_path)
    lines.append(f"DOCX 报告: {output_path}")
    lines.append("完成。")
    write_log(log_path, lines)
    print(output_path)


if __name__ == "__main__":
    main()
